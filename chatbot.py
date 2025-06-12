# import sqlite3
# import os
# import openai
# from dotenv import load_dotenv

# load_dotenv()
# openai.api_key = os.getenv("OPENAI_API_KEY")

# def get_sales_data(query):
#     conn = sqlite3.connect("database.db")
#     cur = conn.cursor()

#     if "today" in query:
#         if "sales" in query:
#             cur.execute("SELECT SUM(amount) FROM sales WHERE date = date('now')")
#             result = cur.fetchone()[0] or 0
#             return f"Today's total sales are ₹{result:.2f}"
#         elif "profit" in query:
#             cur.execute("SELECT SUM(profit) FROM sales WHERE date = date('now')")
#             result = cur.fetchone()[0] or 0
#             return f"Today's total profit is ₹{result:.2f}"
    
#     conn.close()
#     return "Sorry, I couldn't find relevant data."

# def is_company_question(query):
#     keywords = ["sales", "profit", "revenue", "amount", "today", "this week", "earning"]
#     return any(word in query.lower() for word in keywords)

# def ask_bot(user_input):
#     user_input = user_input.lower()

#     if is_company_question(user_input):
#         return get_sales_data(user_input)

#     # Fallback to GPT for general questions
#     try:
#         response = openai.ChatCompletion.create(
#             model="gpt-4o",
#             # model="Gemini API 2.5",  # or "gpt-4" if available
#             messages=[
#                 {"role": "system", "content": "You are a helpful assistant."},
#                 {"role": "user", "content": user_input}
#             ]
#         )
#         return response.choices[0].message["content"]
#     except Exception as e:
#         return "Sorry, I'm currently unable to fetch external information."
    
    
# # def ask_bot(user_input, files=None):
# #     context = ""

# #     # Load content from files
# #     if files:
# #         for file in files:
# #             try:
# #                 with open(file, 'r', encoding='utf-8', errors='ignore') as f:
# #                     context += f"\n---\nContent from {os.path.basename(file)}:\n" + f.read()
# #             except Exception as e:
# #                 context += f"\n---\n[Could not read {file}: {e}]\n"

# #     # Combine user message with context
# #     prompt = f"User Question: {user_input}\n\nRefer to the following context:\n{context}\n\nAnswer:"
    
# #     # Replace this with OpenAI call or your logic
# #     return "This is a placeholder response using context from files."


# import os
# import fitz  # PyMuPDF
# from docx import Document

# def extract_text_from_file(file_path):
#     text = ""
#     ext = os.path.splitext(file_path)[1].lower()

#     try:
#         if ext == ".txt":
#             with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
#                 text = f.read()
#         elif ext == ".pdf":
#             with fitz.open(file_path) as pdf:
#                 for page in pdf:
#                     text += page.get_text()
#         elif ext == ".docx":
#             doc = Document(file_path)
#             for para in doc.paragraphs:
#                 text += para.text + "\n"
#         else:
#             text = f"[Unsupported file type: {ext}]"
#     except Exception as e:
#         text = f"[Error reading file: {file_path}, {e}]"

#     return text

# def ask_bot(user_input, files=None):
#     content_from_files = ""
#     if files:
#         for file_path in files:
#             content_from_files += extract_text_from_file(file_path) + "\n\n"

#     # Basic Q&A logic
#     if "skill" in user_input.lower():
#         if content_from_files:
#             # Extract skills by looking for keyword "Skills" in the file text
#             lines = content_from_files.splitlines()
#             skill_lines = [line for line in lines if "skill" in line.lower()]
#             return "🧠 Extracted skills:\n" + "\n".join(skill_lines) if skill_lines else "⚠️ No skill section found."
#         else:
#             return "📂 Please upload a resume file to extract skills."

#     elif "good morning" in user_input.lower():
#         return "🌞 Good morning! How can I assist you today?"

#     elif "what is in" in user_input.lower() and content_from_files:
#         return "📁 Here is what I found in your uploaded file:\n" + content_from_files[:1000] + "..."

#     return f"🤖 I'm still learning! You said: '{user_input}'"

import os
import openai
import fitz  # PyMuPDF for PDF
from docx import Document
import pandas as pd

# Set your OpenAI API key (recommended to set it via environment variable)
openai.api_key = os.getenv("OPENAI_API_KEY")
# For testing only (not secure): uncomment and paste your API key
# openai.api_key = "your-openai-api-key"

def extract_text_from_file(file_path):
    text = ""
    ext = os.path.splitext(file_path)[1].lower()

    try:
        if ext == ".txt":
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                text = f.read()

        elif ext == ".pdf":
            with fitz.open(file_path) as pdf:
                for page in pdf:
                    text += page.get_text()

        elif ext == ".docx":
            doc = Document(file_path)
            for para in doc.paragraphs:
                text += para.text + "\n"

        elif ext == ".xlsx":
            df = pd.read_excel(file_path)
            text = df.to_string(index=False)

        else:
            text = f"[Unsupported file type: {ext}]"

    except Exception as e:
        text = f"[Error reading file: {file_path}, {e}]"

    return text


def ask_bot(user_input, files=None):
    context = ""

    if files:
        for file_path in files:
            context += extract_text_from_file(file_path) + "\n\n"

    # Truncate to avoid OpenAI token limits
    context = context[:3000]

    # Prompt for GPT
    prompt = f"""
You are a helpful assistant. Use the following context to answer the user's question.

Context:
{context}

User question: {user_input}
Answer:"""

    try:
        response = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": "You are an AI assistant that reads company files and answers questions."},
                {"role": "user", "content": prompt}
            ],
            max_tokens=500,
            temperature=0.4
        )
        return response['choices'][0]['message']['content'].strip()

    except Exception as e:
        return f"⚠️ Error getting response from AI: {e}"
